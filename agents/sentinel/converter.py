"""
Charter & Stone Automation Server
Monitors a local folder for Markdown files, converts them to branded Word documents,
and publishes notifications to Microsoft Teams via Webhook.
"""

import sys
import os
import time
import urllib.parse
from datetime import datetime

# PATH SETUP: Add root to sys.path
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.abspath(os.path.join(current_dir, "../../"))
sys.path.append(project_root)

# Third-party imports
import requests
import markdown
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# Import Shared Auth
from shared.auth import get_graph_headers

# Load environment variables from project root
load_dotenv(os.path.join(project_root, ".env"))

# ================= CONFIGURATION =================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
WATCH_FOLDER = os.path.join(BASE_DIR, "_INBOX")

# --- SMART TEMPLATE FINDER ---
search_paths = [
    os.path.join(BASE_DIR, "Reference.docx"),
    os.path.join(BASE_DIR, "..", "Reference.docx"),
    os.path.join(BASE_DIR, "Reference.docx.docx")
]

TEMPLATE_PATH = None
for path in search_paths:
    full_path = os.path.abspath(path)
    if os.path.exists(full_path):
        TEMPLATE_PATH = full_path
        break

if not TEMPLATE_PATH:
    TEMPLATE_PATH = os.path.join(BASE_DIR, "Reference.docx")

# --- DESTINATION (LOCAL SYNC FOLDER) ---
TARGET_ROOT = os.getenv("TARGET_ROOT")

# --- SHAREPOINT WEB URL ---
SHAREPOINT_FOLDER_URL = os.getenv("SHAREPOINT_FOLDER_URL")

# --- TEAMS WORKFLOW URL ---
TEAMS_WEBHOOK_URL = os.getenv("TEAMS_WEBHOOK_URL")

# =================================================

def notify_teams(filename, _filepath):
    """Sends a WRAPPED Adaptive Card with a Clickable Link."""
    if "PASTE_YOUR" in TEAMS_WEBHOOK_URL:
        print("⚠️ Skipping Teams Alert (Webhook URL not set).")
        return

    # Generate Smart Link
    if "charterandstone" in SHAREPOINT_FOLDER_URL:
        safe_filename = urllib.parse.quote(filename)
        base = SHAREPOINT_FOLDER_URL.rstrip('/')
        # ADDing '?web=1' forces SharePoint to open the Word Online viewer
        web_link = f"{base}/{safe_filename}?web=1"
        display_value = f"[{filename}]({web_link})"
    else:
        display_value = filename

    wrapped_card_payload = {
        "type": "message",
        "attachments": [
            {
                "contentType": "application/vnd.microsoft.card.adaptive",
                "content": {
                    "type": "AdaptiveCard",
                    "version": "1.4",
                    "body": [
                        {
                            "type": "TextBlock",
                            "text": "🚀 New Strategy Asset Published",
                            "weight": "Bolder",
                            "size": "Medium",
                            "color": "Accent"
                        },
                        {
                            "type": "TextBlock",
                            "text": "Charter & Stone Autonomous Agent",
                            "isSubtle": True,
                            "wrap": True,
                            "size": "Small"
                        },
                        {
                            "type": "FactSet",
                            "facts": [
                                {"title": "Document:", "value": display_value},
                                {"title": "Location:", "value": "Strategy & Intel"},
                                {"title": "Status:", "value": "Synced & Ready"}
                            ]
                        }
                    ],
                    "msteams": {
                        "width": "Full"
                    }
                }
            }
        ]
    }

    try:
        response = requests.post(
            TEAMS_WEBHOOK_URL, 
            json=wrapped_card_payload,
            headers={'Content-Type': 'application/json'},
            timeout=10
        )
        if response.status_code == 202: 
            print("📣 Teams Alert Sent (202 Accepted)!")
        elif response.status_code == 200:
            print("📣 Teams Alert Sent (200 OK)!")
        else:
            print(f"⚠️ Teams Alert Failed: {response.status_code} - {response.text}")
    except Exception as e:
        print(f"⚠️ Error sending alert: {e}")


def apply_branding_to_doc(doc, title_text):
    """Applies the Charter & Stone header and date styles."""
    doc.add_paragraph("") 
    head = doc.add_paragraph(title_text, style='Heading 1')
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("") 


def convert_md_to_branded_docx(md_file_path):
    print(f"⚙️ Processing: {md_file_path}")

    try:
        if os.path.exists(TEMPLATE_PATH):
            doc = Document(TEMPLATE_PATH)
            # CLEAR SCAFFOLDING
            for paragraph in doc.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
            print(f"🎨 Branding loaded from: {TEMPLATE_PATH}")
        else:
            doc = Document()
            print("⚠️ Template not found. Using default styles.")

        with open(md_file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')

        base_name = os.path.basename(md_file_path).replace('.md', '')
        display_title = base_name.replace('_', ' ').title()
        apply_branding_to_doc(doc, display_title)

        for element in soup.contents:
            if not isinstance(element, Tag):
                continue
            
            if element.name == 'h1':
                doc.add_heading(element.text, level=1) 
            elif element.name == 'h2':
                doc.add_heading(element.text, level=2)
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Number')

        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        filename = f"Report_{base_name}_{timestamp}.docx"

        if not os.path.exists(TARGET_ROOT):
            os.makedirs(TARGET_ROOT)

        final_path = os.path.join(TARGET_ROOT, filename)
        doc.save(final_path)

        print(f"✅ PUBLISHED: {final_path}")
        notify_teams(filename, final_path)
        os.rename(md_file_path, md_file_path + ".processed")

    except Exception as error:
        print(f"❌ Error: {error}")


def process_inbox():
    """Process all markdown files in the inbox folder (called by daemon)"""
    if not os.path.exists(WATCH_FOLDER):
        os.makedirs(WATCH_FOLDER)
    
    files_found = [f for f in os.listdir(WATCH_FOLDER) if f.endswith(".md")]
    if files_found:
        print(f"📂 Found {len(files_found)} file(s) to process")
        for file in files_found:
            full_path = os.path.join(WATCH_FOLDER, file)
            convert_md_to_branded_docx(full_path)
    else:
        print("📭 Inbox empty")


if __name__ == "__main__":
    if not os.path.exists(WATCH_FOLDER):
        os.makedirs(WATCH_FOLDER)
    print("--- CHARTER & STONE AUTOMATION SERVER ONLINE ---")
    print(f"👀 Watching: {WATCH_FOLDER}")
    print(f"🎯 Target: {TARGET_ROOT}")

    while True:
        if os.path.exists(WATCH_FOLDER):
            files_found = [f for f in os.listdir(WATCH_FOLDER) if f.endswith(".md")]
            for file in files_found:
                full_path = os.path.join(WATCH_FOLDER, file)
                convert_md_to_branded_docx(full_path)
        time.sleep(3)