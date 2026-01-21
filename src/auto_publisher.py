import os
import time
import json
import requests
import markdown
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ================= CONFIGURATION =================

# Get the folder where this script is actually living right now (src)
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
# Get the project root (one level up)
BASE_DIR = os.path.dirname(SCRIPT_DIR)

WATCH_FOLDER = os.path.join(BASE_DIR, "_INBOX")

# --- SMART TEMPLATE FINDER ---
possible_names = ["Reference.docx", "Reference.docx.docx"]
TEMPLATE_PATH = None
for name in possible_names:
    path = os.path.join(BASE_DIR, name)
    if os.path.exists(path):
        TEMPLATE_PATH = path
        break
if not TEMPLATE_PATH:
    TEMPLATE_PATH = os.path.join(BASE_DIR, "Reference.docx")

# --- DESTINATION ---
TARGET_ROOT = r"C:\Users\tasms\Charter & Stone\Charter & Stone HQ - Documents\02_Research_Intel"

# --- TEAMS WORKFLOW URL ---
# PASTE YOUR LONG URL INSIDE THE QUOTES BELOW
TEAMS_WEBHOOK_URL = "https://defaulta7704ac7417c49ecb21418a1636bd1.cd.environment.api.powerplatform.com:443/powerautomate/automations/direct/workflows/fda92e1484ac4a6dbfb48f4805e098b8/triggers/manual/paths/invoke?api-version=1&sp=%2Ftriggers%2Fmanual%2Frun&sv=1.0&sig=3a1FWfN885HeC5_Z4Za_sIqbp86WDk_D4xz5nTCqmqU"

# =================================================

def notify_teams(filename, filepath):
    """Sends a modern Adaptive Card to the Teams Channel."""
    if "PASTE_YOUR" in TEAMS_WEBHOOK_URL:
        print("⚠️ Skipping Teams Alert (Webhook URL not set).")
        return

    # THE FIX: Sending raw Adaptive Card JSON without the 'message' wrapper.
    card_payload = {
        "type": "AdaptiveCard",
        "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
        "version": "1.4",
        "body": [
            {
                "type": "TextBlock",
                "text": "🚀 New Strategy Asset Published",
                "weight": "Bolder",
                "size": "Medium",
                "color": "Accent"
            },
            {
                "type": "TextBlock",
                "text": "Charter & Stone Autonomous Agent",
                "isSubtle": True,
                "wrap": True,
                "size": "Small"
            },
            {
                "type": "FactSet",
                "facts": [
                    {"title": "Document:", "value": filename},
                    {"title": "Location:", "value": "02_Research_Intel"},
                    {"title": "Status:", "value": "Synced & Ready"}
                ]
            }
        ],
        "msteams": {
            "width": "Full"
        }
    }

    try:
        response = requests.post(
            TEAMS_WEBHOOK_URL, 
            json=card_payload,
            headers={'Content-Type': 'application/json'}
        )
        if response.status_code == 202: 
            print("📣 Teams Alert Sent (Workflow Triggered)!")
        elif response.status_code == 200:
            print("📣 Teams Alert Sent!")
        else:
            print(f"⚠️ Teams Alert Failed: {response.status_code} - {response.text}")
    except Exception as e:
        print(f"⚠️ Error sending alert: {e}")

def apply_branding_to_doc(doc, title_text):
    doc.add_paragraph("") 
    head = doc.add_paragraph(title_text, style='Heading 1')
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("") 

def convert_md_to_branded_docx(md_file_path):
    print(f"⚙️ Processing: {md_file_path}")

    try:
        if os.path.exists(TEMPLATE_PATH):
            doc = Document(TEMPLATE_PATH)
        else:
            doc = Document()

        with open(md_file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')

        base_name = os.path.basename(md_file_path).replace('.md', '')
        display_title = base_name.replace('_', ' ').title()
        apply_branding_to_doc(doc, display_title)

        for element in soup.contents:
            if not isinstance(element, Tag): continue
            if element.name == 'h1': doc.add_heading(element.text, level=2) 
            elif element.name == 'h2': doc.add_heading(element.text, level=3)
            elif element.name == 'p': doc.add_paragraph(element.text)
            elif element.name == 'ul':
                for li in element.find_all('li'): doc.add_paragraph(li.text, style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li'): doc.add_paragraph(li.text, style='List Number')

        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        filename = f"Report_{base_name}_{timestamp}.docx"
        
        if not os.path.exists(TARGET_ROOT): os.makedirs(TARGET_ROOT)  
        final_path = os.path.join(TARGET_ROOT, filename)
        doc.save(final_path)
        
        print(f"✅ PUBLISHED: {final_path}")
        
        # --- TRIGGER NOTIFICATION ---
        notify_teams(filename, final_path)

        processed_path = md_file_path + ".processed"
        if os.path.exists(processed_path):
            os.remove(processed_path)
        os.rename(md_file_path, processed_path)

    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    if not os.path.exists(WATCH_FOLDER): os.makedirs(WATCH_FOLDER)
    print("--- CHARTER & STONE AUTOMATION SERVER ONLINE ---")
    print(f"👀 Watching: {WATCH_FOLDER}")
    
    while True:
        if os.path.exists(WATCH_FOLDER):
            files_found = [f for f in os.listdir(WATCH_FOLDER) if f.endswith(".md")]
            for file in files_found:
                full_path = os.path.join(WATCH_FOLDER, file)
                convert_md_to_branded_docx(full_path)
        time.sleep(3)